from flask import Blueprint, request, jsonify, current_app, send_from_directory, send_file
from werkzeug.utils import secure_filename
from datetime import datetime
import os

from extensions import db
from models import FreaderContract, CutAIContract, FreaderContractVersion, CutAIContractVersion
from service.db_service import save_contract_to_db
from service.ai_service import analyze_contract_text
from utils.parser import parse_document

contracts_bp = Blueprint('contracts', __name__)


@contracts_bp.route('/upload-and-analyze', methods=['POST'])
def upload_and_analyze():
    # 1. Verifica file
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "Nessun file fornito"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"status": "error", "message": "Nome file vuoto"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

    file.save(filepath)

    try:
        # 3. Estrazione testo (OCR o DOCX)
        print(f"📄 Estrazione testo da {filename}...")
        text = parse_document(filepath)
        
        if not text or len(text.strip()) < 10:
            return jsonify({"status": "error", "message": "Impossibile estrarre testo dal documento"}), 422

        # 4. Analisi con LLM
        print(f"🧠 Analisi IA in corso...")
        analysis_result = analyze_contract_text(text)

        # 5. Ritorna il JSON all'utente per la verifica (include filename)
        return jsonify({
            "status": "success",
            "data": {
                "filename": filename,
                "analysis": analysis_result
            }
        }), 200

    except Exception as e:
        print(f"❌ Errore durante il processo: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 0. BATCH UPLOAD AND ANALYZE (POST)
# ==========================================
@contracts_bp.route('/upload-and-analyze-batch', methods=['POST'])
def upload_and_analyze_batch():
    if 'files' not in request.files:
        return jsonify({"status": "error", "message": "Nessun file fornito (chiave 'files' mancante)"}), 400
        
    files = request.files.getlist('files')
    if not files or all(file.filename == '' for file in files):
        return jsonify({"status": "error", "message": "Nessun file valido trovato"}), 400

    results = []
    errors = []

    for file in files:
        if file.filename == '':
            continue
            
        filename = secure_filename(file.filename)
        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

        file.save(filepath)

        try:
            print(f"📄 Estrazione testo da {filename}...")
            text = parse_document(filepath)
            
            if not text or len(text.strip()) < 10:
                print(f"⚠️ Testo non sufficiente per {filename}")
                errors.append({"filename": filename, "error": "Impossibile estrarre testo o testo troppo breve"})
                continue

            print(f"🧠 Analisi IA in corso per {filename}...")
            analysis_result = analyze_contract_text(text)
            
            results.append({
                "filename": filename,
                "analysis": analysis_result
            })
            
        except Exception as e:
            print(f"❌ Errore durante l'elaborazione di {filename}: {str(e)}")
            errors.append({"filename": filename, "error": str(e)})

    return jsonify({
        "status": "success",
        "data": results,
        "errors": errors
    }), 200

# ==========================================
# 1. UPLOAD (POST)
# ==========================================
@contracts_bp.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "Nessun file fornito"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"status": "error", "message": "Nome file vuoto"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

    file.save(filepath)

    try:
        # Estrazione testo grezzo
        extracted_text = parse_document(filepath)
        return jsonify({
            "status": "success",
            "data": {
                "filename": filename,
                "extracted_text": extracted_text
            }
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 2. ANALYZE (POST)
# ==========================================
@contracts_bp.route('/analyze', methods=['POST'])
def analyze_text():
    data = request.get_json()
    text = data.get('text')
    
    if not text:
        return jsonify({"status": "error", "message": "Testo mancante"}), 400

    try:
        struttura_json = analyze_contract_text(text)
        return jsonify({
            "status": "success",
            "data": struttura_json
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 3. SAVE (POST) - Conferma Utente
# ==========================================
@contracts_bp.route('/save', methods=['POST'])
def save_contract():
    tenant_id = request.headers.get('X-Tenant-ID')
    data = request.get_json()
    user = data.get('user_id', 'System')
    extracted_data = data.get('extracted_data')
    filename = data.get('filename')

    if not extracted_data:
        return jsonify({"status": "error", "message": "Dati estratti mancanti"}), 400

    try:
        contract_id = save_contract_to_db(tenant_id, extracted_data, created_by=user, filename=filename)
        return jsonify({"status": "success", "contract_id": contract_id}), 201
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 4. LIST (GET) - Data Table Frontend
# ==========================================
@contracts_bp.route('/list', methods=['GET'])
def list_contracts():
    tenant_id = request.headers.get('X-Tenant-ID')
    unified_list = []

    # Recuperiamo e uniformiamo i contratti Freader
    freader_contracts = FreaderContract.query.filter_by(tenant_id=tenant_id).all()
    for c in freader_contracts:
        unified_list.append({
            "id": c.id,
            "prodotto": "Freader",
            "cliente": c.cliente_ragione_sociale,
            "status": c.status,
            "versioni": len(c.versions)
        })

    # Recuperiamo e uniformiamo i contratti CutAI
    cutai_contracts = CutAIContract.query.filter_by(tenant_id=tenant_id).all()
    for c in cutai_contracts:
        unified_list.append({
            "id": c.id,
            "prodotto": "CutAI",
            "cliente": c.cliente_ragione_sociale,
            "status": c.status,
            "versioni": len(c.versions)
        })

    return jsonify({"status": "success", "data": unified_list}), 200

# ==========================================
# 4b. MAP DATA (GET) - Contratti con coordinate
# ==========================================
@contracts_bp.route('/map', methods=['GET'])
def get_map_contracts():
    tenant_id = request.headers.get('X-Tenant-ID')
    features = []

    freader_contracts = FreaderContract.query.filter_by(tenant_id=tenant_id).all()
    for c in freader_contracts:
        if c.lat and c.lng:
            latest = c.versions[0] if c.versions else None
            features.append({
                "type": "Feature",
                "geometry": {"type": "Point", "coordinates": [c.lng, c.lat]},
                "properties": {
                    "id": c.id,
                    "prodotto": "Freader",
                    "cliente": c.cliente_ragione_sociale,
                    "sede": c.cliente_sede_legale,
                    "status": c.status,
                    "canone": latest.canone_trimestrale if latest else None
                }
            })

    cutai_contracts = CutAIContract.query.filter_by(tenant_id=tenant_id).all()
    for c in cutai_contracts:
        if c.lat and c.lng:
            latest = c.versions[0] if c.versions else None
            features.append({
                "type": "Feature",
                "geometry": {"type": "Point", "coordinates": [c.lng, c.lat]},
                "properties": {
                    "id": c.id,
                    "prodotto": "CutAI",
                    "cliente": c.cliente_ragione_sociale,
                    "sede": c.cliente_sede_legale,
                    "status": c.status,
                    "canone": latest.canone_base_trimestrale if latest else None
                }
            })

    return jsonify({
        "type": "FeatureCollection",
        "features": features
    }), 200

# ==========================================
# 5. DOWNLOAD PDF/FILE (GET) - PDF Viewer
# ==========================================
@contracts_bp.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Permette al frontend di visualizzare il PDF. 
    Esempio in React: <iframe src="http://localhost:5000/api/v1/contracts/download/contratto.pdf" />
    """
    return send_from_directory(current_app.config['UPLOAD_FOLDER'], secure_filename(filename))

# ==========================================
# 6. GET SINGOLO CONTRATTO (GET)
# ==========================================
@contracts_bp.route('/<contract_id>', methods=['GET'])
def get_contract(contract_id):
    tenant_id = request.headers.get('X-Tenant-ID')
    
    # Cerchiamo in Freader
    f_contract = FreaderContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if f_contract:
        return jsonify({
            "status": "success",
            "data": {
                "id": f_contract.id,
                "prodotto": "Freader",
                "cliente": f_contract.cliente_ragione_sociale,
                "sede": f_contract.cliente_sede_legale,
                "filename": f_contract.original_filename,
                "history": [{"version": v.version_number, "data_firma": v.data_firma, "canone": v.canone_trimestrale} for v in f_contract.versions]
            }
        }), 200

    # Cerchiamo in CutAI
    c_contract = CutAIContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if c_contract:
        return jsonify({
            "status": "success",
            "data": {
                "id": c_contract.id,
                "prodotto": "CutAI",
                "cliente": c_contract.cliente_ragione_sociale,
                "sede": c_contract.cliente_sede_legale,
                "filename": c_contract.original_filename,
                "history": [{"version": v.version_number, "piano": v.profilo_commerciale, "canone": v.canone_base_trimestrale} for v in c_contract.versions]
            }
        }), 200

    return jsonify({"status": "error", "code": 404, "message": "Contratto non trovato."}), 404

# ==========================================
# 7. CREA NUOVA VERSIONE (POST / UPDATE)
# ==========================================
@contracts_bp.route('/<contract_id>/versions', methods=['POST'])
def add_new_version(contract_id):
    tenant_id = request.headers.get('X-Tenant-ID')
    data = request.get_json()
    
    # Verifica se è Freader
    f_contract = FreaderContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if f_contract:
        last_version = FreaderContractVersion.query.filter_by(contract_id=contract_id).order_by(FreaderContractVersion.version_number.desc()).first()
        new_version_num = (last_version.version_number + 1) if last_version else 1
        
        # Inserisci nuova riga (Audit-Ready)
        new_v = FreaderContractVersion(
            contract_id=contract_id,
            version_number=new_version_num,
            data_firma=data.get('data_firma', last_version.data_firma if last_version else ''),
            durata_mesi=data.get('durata_mesi', last_version.durata_mesi if last_version else 12),
            preavviso_giorni=data.get('preavviso_giorni', last_version.preavviso_giorni if last_version else 30),
            canone_trimestrale=data.get('canone_trimestrale', last_version.canone_trimestrale if last_version else 0),
            prezzo_fascia_1=data.get('prezzo_fascia_1', last_version.prezzo_fascia_1 if last_version else 0),
            prezzo_fascia_2=data.get('prezzo_fascia_2', last_version.prezzo_fascia_2 if last_version else 0),
            prezzo_fascia_3=data.get('prezzo_fascia_3', last_version.prezzo_fascia_3 if last_version else 0),
            credito_uptime=data.get('credito_uptime', last_version.credito_uptime if last_version else 0),
            credito_ticketing=data.get('credito_ticketing', last_version.credito_ticketing if last_version else 0),
            tetto_crediti=data.get('tetto_crediti', last_version.tetto_crediti if last_version else 0),
            created_by=data.get('user_id', 'System'),
            change_reason=data.get('change_reason', 'Rinegoziazione')
        )
        db.session.add(new_v)
        db.session.commit()
        return jsonify({"status": "success", "new_version": new_version_num}), 201

    # Verifica se è CutAI
    c_contract = CutAIContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if c_contract:
        last_version = CutAIContractVersion.query.filter_by(contract_id=contract_id).order_by(CutAIContractVersion.version_number.desc()).first()
        new_version_num = (last_version.version_number + 1) if last_version else 1
        
        # Inserisci nuova riga (Audit-Ready)
        new_v = CutAIContractVersion(
            contract_id=contract_id,
            version_number=new_version_num,
            data_sottoscrizione=data.get('data_firma', last_version.data_sottoscrizione if last_version else ''),
            durata_mesi=data.get('durata_mesi', last_version.durata_mesi if last_version else 12),
            preavviso_giorni=data.get('preavviso_giorni', last_version.preavviso_giorni if last_version else 30),
            profilo_commerciale=data.get('profilo_commerciale', last_version.profilo_commerciale if last_version else 'Standard'),
            canone_base_trimestrale=data.get('canone_trimestrale', last_version.canone_base_trimestrale if last_version else 0),
            soglia_utenti_inclusi=data.get('soglia_utenti_inclusi', last_version.soglia_utenti_inclusi if last_version else 0),
            fee_utente_extra=data.get('fee_utente_extra', last_version.fee_utente_extra if last_version else 0),
            soglia_minima_servizio=data.get('soglia_minima_servizio', last_version.soglia_minima_servizio if last_version else 98.0),
            credito_uptime=data.get('credito_uptime', last_version.credito_uptime if last_version else 0),
            credito_ticketing=data.get('credito_ticketing', last_version.credito_ticketing if last_version else 0),
            tetto_crediti=data.get('tetto_crediti', last_version.tetto_crediti if last_version else 0),
            created_by=data.get('user_id', 'System'),
            change_reason=data.get('change_reason', 'Rinegoziazione')
        )
        db.session.add(new_v)
        db.session.commit()
        return jsonify({"status": "success", "new_version": new_version_num}), 201
    
    return jsonify({"status": "error", "code": 404, "message": "Contratto non trovato."}), 404


# ==========================================
# 8. PRE-SIGN ANALYSIS (POST)
# ==========================================
@contracts_bp.route('/pre-sign-analysis', methods=['POST'])
def pre_sign_analysis():
    """
    Analizza un contratto pre-firma confrontandolo con lo storico.
    Restituisce 3 livelli di modifiche suggerite.
    """
    from service.ai_service import analyze_pre_sign
    from datetime import datetime, timedelta
    
    tenant_id = request.headers.get('X-Tenant-ID')
    data = request.get_json()
    
    contract_data = data.get('contract_data')
    if not contract_data:
        return jsonify({"status": "error", "message": "Dati contratto mancanti"}), 400
    
    prodotto = contract_data.get('prodotto')
    
    try:
        # Get historical contracts of same product
        historical = []
        today = datetime.now()
        
        if prodotto == 'Freader':
            contracts = FreaderContract.query.filter_by(tenant_id=tenant_id, status='ACTIVE').all()
            for c in contracts:
                if c.versions:
                    latest = c.versions[0]
                    try:
                        data_firma = datetime.strptime(latest.data_firma, '%d/%m/%Y')
                        scadenza = data_firma + timedelta(days=latest.durata_mesi * 30)
                        giorni_scadenza = (scadenza - today).days
                        
                        historical.append({
                            'cliente': c.cliente_ragione_sociale,
                            'canone_trim': latest.canone_trimestrale,
                            'durata_mesi': latest.durata_mesi,
                            'preavviso_gg': latest.preavviso_giorni,
                            'tetto_cred': latest.tetto_crediti,
                            'credito_uptime': latest.credito_uptime,
                            'credito_ticketing': latest.credito_ticketing,
                            'giorni_scadenza': giorni_scadenza
                        })
                    except ValueError:
                        pass
        
        elif prodotto == 'CutAI':
            contracts = CutAIContract.query.filter_by(tenant_id=tenant_id, status='ACTIVE').all()
            for c in contracts:
                if c.versions:
                    latest = c.versions[0]
                    try:
                        data_firma = datetime.strptime(latest.data_sottoscrizione, '%d/%m/%Y')
                        scadenza = data_firma + timedelta(days=latest.durata_mesi * 30)
                        giorni_scadenza = (scadenza - today).days
                        
                        historical.append({
                            'cliente': c.cliente_ragione_sociale,
                            'canone_trim': latest.canone_base_trimestrale,
                            'durata_mesi': latest.durata_mesi,
                            'preavviso_gg': latest.preavviso_giorni,
                            'tetto_cred': latest.tetto_crediti,
                            'credito_uptime': latest.credito_uptime,
                            'credito_ticketing': latest.credito_ticketing,
                            'giorni_scadenza': giorni_scadenza
                        })
                    except ValueError:
                        pass
        
        # Run AI analysis
        analysis = analyze_pre_sign(contract_data, historical)
        
        return jsonify({
            "status": "success",
            "data": analysis
        }), 200
        
    except Exception as e:
        print(f"❌ Errore analisi pre-firma: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 9. RISK ANALYSIS (POST)
# ==========================================
@contracts_bp.route('/risk-analysis', methods=['POST'])
def risk_analysis():
    """Analizza rischio contratto: risk score, punti critici, errori ortografici."""
    from service.ai_service import analyze_contract_risk
    
    data = request.get_json()
    contract_data = data.get('contract_data')
    testo_originale = data.get('testo_originale', '')
    
    if not contract_data:
        return jsonify({"status": "error", "message": "Dati contratto mancanti"}), 400
    
    try:
        result = analyze_contract_risk(contract_data, testo_originale)
        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:
        print(f"❌ Errore risk analysis: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 10. IMPROVE CLAUSE (POST)
# ==========================================
@contracts_bp.route('/improve-clause', methods=['POST'])
def improve_clause():
    """AI riscrive una clausola critica con termini migliori."""
    from service.ai_service import improve_contract_clause
    
    data = request.get_json()
    clausola_originale = data.get('clausola_originale', '')
    tipo_problema = data.get('tipo_problema', '')
    contesto = data.get('contesto_contratto', '')
    valore_rif = data.get('valore_riferimento', '')
    
    if not clausola_originale:
        return jsonify({"status": "error", "message": "Clausola originale mancante"}), 400
    
    try:
        result = improve_contract_clause(clausola_originale, tipo_problema, contesto, valore_rif)
        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:
        print(f"❌ Errore improve clause: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 11. RECALCULATE RISK (POST)
# ==========================================
@contracts_bp.route('/recalculate-risk', methods=['POST'])
def recalculate_risk():
    """Ricalcola risk score con le modifiche applicate."""
    from service.ai_service import analyze_contract_risk
    
    data = request.get_json()
    contract_data = data.get('contract_data')
    
    if not contract_data:
        return jsonify({"status": "error", "message": "Dati contratto mancanti"}), 400
    
    try:
        result = analyze_contract_risk(contract_data, '')
        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 12. DOWNLOAD MODIFIED CONTRACT DOCX (POST)
# ==========================================
@contracts_bp.route('/download-modified-pdf', methods=['POST'])
def download_modified_docx():
    """Apre il DOCX originale caricato, aggiunge in fondo le sezioni modificate."""
    from io import BytesIO
    from flask import send_file
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    data = request.get_json()
    contract_data = data.get('contract_data', {})
    original_risk = data.get('original_risk', {})
    modifications = data.get('modifications', {})
    new_risk_score = data.get('new_risk_score', 0)
    filename = data.get('filename', '')
    
    ana = contract_data.get('anagrafica', {})
    punti = original_risk.get('punti_critici', [])
    orig_score = original_risk.get('risk_score', 0)
    modified_ids = set(modifications.keys()) if modifications else set()
    
    # Try to open the original uploaded file
    original_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename) if filename else None
    
    if original_path and os.path.exists(original_path) and original_path.endswith('.docx'):
        doc = Document(original_path)
    else:
        # Fallback: create new doc with contract data
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        prodotto = contract_data.get('prodotto', 'N/A')
        det = contract_data.get('dettagli_contratto', {})
        sla = contract_data.get('sla', {})
        
        doc.add_heading(f'Contratto — {ana.get("cliente_ragione_sociale", "N/A")}', level=0)
        p = doc.add_paragraph()
        r = p.add_run(f'Prodotto: {prodotto} | Data: {det.get("data_firma", "N/A")}')
        r.font.size = Pt(10)
        
        doc.add_heading('Anagrafica', level=2)
        doc.add_paragraph(f'Ragione Sociale: {ana.get("cliente_ragione_sociale", "N/A")}')
        doc.add_paragraph(f'Sede Legale: {ana.get("cliente_sede_legale", "N/A")}')
        
        doc.add_heading('Dettagli Contratto', level=2)
        doc.add_paragraph(f'Data Firma: {det.get("data_firma", "N/A")}')
        doc.add_paragraph(f'Durata: {det.get("durata_mesi", "N/A")} mesi')
        doc.add_paragraph(f'Preavviso: {det.get("preavviso_giorni", "N/A")} giorni')
        
        doc.add_heading('SLA', level=2)
        doc.add_paragraph(f'Credito Uptime: {sla.get("credito_uptime", "N/A")}%')
        doc.add_paragraph(f'Credito Ticketing: {sla.get("credito_ticketing", "N/A")}%')
        doc.add_paragraph(f'Tetto Crediti: {sla.get("tetto_crediti", "N/A")}%')
    
    # ── APPEND: Modifications section at the end ──
    if modified_ids:
        doc.add_page_break()
        
        h = doc.add_heading('ALLEGATO — Modifiche Proposte da Nexus Core AI', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        
        # Risk score table
        p_score = doc.add_paragraph()
        r = p_score.add_run(f'Risk Score: {orig_score}% → {new_risk_score}% (delta {new_risk_score - orig_score:+d}%)')
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x05, 0x96, 0x69) if new_risk_score < orig_score else RGBColor(0xDC, 0x26, 0x26)
        
        doc.add_paragraph()
        
        for p_item in punti:
            if p_item.get('id') not in modified_ids:
                continue
            mod = modifications[p_item['id']]
            
            # Section heading
            h2 = doc.add_heading(p_item.get('sezione', ''), level=2)
            for run in h2.runs:
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
            # Severity
            p_sev = doc.add_paragraph()
            r_sev = p_sev.add_run(f'Gravità: {p_item.get("gravita", "").upper()}')
            r_sev.font.size = Pt(9)
            r_sev.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) if p_item.get('gravita') == 'alta' else RGBColor(0xD9, 0x77, 0x06)
            
            # Original text (strikethrough)
            p_orig = doc.add_paragraph()
            r_label = p_orig.add_run('ORIGINALE: ')
            r_label.bold = True
            r_label.font.size = Pt(9)
            r_label.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            r_text = p_orig.add_run(p_item.get('testo_contratto_originale', ''))
            r_text.font.size = Pt(10)
            r_text.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            r_text.font.strike = True
            
            # New text (green, bold)
            p_new = doc.add_paragraph()
            r_label2 = p_new.add_run('NUOVO: ')
            r_label2.bold = True
            r_label2.font.size = Pt(9)
            r_label2.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            r_new = p_new.add_run(mod.get('testo_migliorato', ''))
            r_new.font.size = Pt(10)
            r_new.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            r_new.bold = True
            
            if mod.get('motivazione'):
                p_mot = doc.add_paragraph()
                r_mot = p_mot.add_run(f'Motivazione: {mod["motivazione"]}')
                r_mot.font.size = Pt(8)
                r_mot.font.italic = True
                r_mot.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            
            doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = footer.add_run(f'Generato da Nexus Core Contract Intelligence — {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    out_name = f'contratto_{ana.get("cliente_ragione_sociale", "contratto").replace(" ", "_")}.docx'
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=out_name)
